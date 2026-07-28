# -*- coding: utf-8 -*-
"""Cadeia canônica e auditável de extração do POP da Academia IAT.

O script:
1. valida a identidade da minuta-fonte v1.7 por nome, tamanho e SHA-256;
2. preserva os IDs de seção existentes;
3. extrai texto, tabelas, figuras e somente metadados públicos necessários;
4. atualiza o catálogo, o manifesto de ativos e a validação de fidelidade;
5. só publica os artefatos se todos os gates passarem.

O material permanece uma minuta técnica, pendente de validação humana e
institucional. Ele não é tratado por esta rotina como norma ou ato oficial.

Uso:
    python tools/extract_pop.py "C:\\caminho\\para\\o POP.docx"
    IAT_POP_SOURCE="C:\\caminho\\para\\o POP.docx" python tools/extract_pop.py
"""
from __future__ import annotations

import argparse
import hashlib
import io
import json
import mimetypes
import os
import re
import sys
import unicodedata
import zipfile
from datetime import datetime, timezone
from pathlib import Path
from xml.etree import ElementTree as ET

import docx
from docx.document import Document as DocxDocument
from docx.oxml.ns import qn
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

ROOT = Path(__file__).resolve().parents[1]
DATA_DIR = ROOT / "src" / "data"
OUT_JSON = DATA_DIR / "pop-content.json"
FLOW_JSON = DATA_DIR / "flowcharts-content.json"
CATALOG_JSON = DATA_DIR / "content-catalog.json"
VALIDATION_JSON = DATA_DIR / "extraction-validation.json"
ASSET_DIR = ROOT / "public" / "source-assets"
ASSET_MANIFEST = ASSET_DIR / "asset-manifest.json"

DEFAULT_POP_SOURCE = (
    Path(os.environ["IAT_POP_SOURCE"]) if os.environ.get("IAT_POP_SOURCE") else None
)

PIPELINE_VERSION = "2.0.0"
EXPECTED = {
    "fileName": "POP ou Manual Hidreletricas IAT Julho de 2026 (Com APA, UCs, RTTA).docx",
    "bytes": 4_408_377,
    "sha256": "8ffa771546c244e194e6d7b41dd91d5ab3f56083e94c081e1e5c9a17f13f2c3c",
    "version": "1.7",
    "sections": 167,
    "learningSections": 161,
    "navigationSections": 6,
    "tables": 66,
    "quadros": 46,
    "tabelas": 20,
    "figures": 14,
    "assets": 14,
    "paragraphNodes": 3_339,
    "bodyBlocks": 765,
}

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NAVIGATION_ROOTS = {
    "sumario navegavel",
    "indice de fluxogramas",
    "indice navegavel de quadros e tabelas",
    "indice de anexos",
}

PUBLIC_CORE_PROPERTIES = {
    "title",
    "subject",
    "keywords",
    "description",
    "revision",
    "created",
    "modified",
    "category",
    "contentStatus",
    "contentType",
    "identifier",
    "language",
    "version",
}
PUBLIC_APPLICATION_PROPERTIES = {
    "Template",
    "TotalTime",
    "Pages",
    "Words",
    "Characters",
    "Application",
    "DocSecurity",
    "Lines",
    "Paragraphs",
    "ScaleCrop",
    "HeadingPairs",
    "TitlesOfParts",
    "LinksUpToDate",
    "CharactersWithSpaces",
    "SharedDoc",
    "HyperlinksChanged",
    "AppVersion",
}


def norm(text: str) -> str:
    text = unicodedata.normalize("NFD", text or "")
    text = "".join(char for char in text if unicodedata.category(char) != "Mn")
    return re.sub(r"[^a-z0-9]+", " ", text.lower()).strip()


def normalized_fidelity_text(text: str) -> str:
    return re.sub(r"[\s\u00ad\u2011]", "", text or "")


def sha256_bytes(payload: bytes) -> str:
    return hashlib.sha256(payload).hexdigest()


def iso_utc(timestamp: float) -> str:
    value = datetime.fromtimestamp(timestamp, timezone.utc)
    return value.isoformat(timespec="milliseconds").replace("+00:00", "Z")


def natural_key(value: str):
    return [int(part) if part.isdigit() else part.lower() for part in re.split(r"(\d+)", value)]


def local_name(tag: str) -> str:
    return tag.rsplit("}", 1)[-1]


def read_zip_properties(archive: zipfile.ZipFile, part_name: str) -> dict:
    try:
        root = ET.fromstring(archive.read(part_name))
    except KeyError:
        return {}
    result: dict[str, object] = {}
    for child in root:
        key = local_name(child.tag)
        values = [(node.text or "") for node in child.iter() if node is not child and node.text]
        value: object = (child.text or "").strip() if not values else (values[0] if len(values) == 1 else values)
        if key in result:
            prior = result[key]
            result[key] = [*prior, value] if isinstance(prior, list) else [prior, value]
        else:
            result[key] = value
    return result


def public_properties(properties: dict, allowlist: set[str]) -> dict:
    """Mantém metadados técnicos úteis e descarta autoria/identificadores pessoais."""
    return {key: value for key, value in properties.items() if key in allowlist}


def personal_metadata_values(core: dict, application: dict) -> set[str]:
    candidates = [
        core.get("creator"),
        core.get("lastModifiedBy"),
        application.get("Manager"),
    ]
    return {
        value.strip()
        for value in candidates
        if isinstance(value, str) and len(value.strip()) >= 3
    }


def redact_public_value(value, personal_values: set[str]) -> tuple[object, int]:
    """Remove valores pessoais derivados dos metadados antes da publicação."""
    if isinstance(value, str):
        result = value
        count = 0
        for personal_value in sorted(personal_values, key=len, reverse=True):
            result, replacements = re.subn(
                re.escape(personal_value),
                "[nome removido por privacidade]",
                result,
                flags=re.IGNORECASE,
            )
            count += replacements
        return result, count
    if isinstance(value, list):
        redacted_list = []
        count = 0
        for item in value:
            redacted_item, replacements = redact_public_value(item, personal_values)
            redacted_list.append(redacted_item)
            count += replacements
        return redacted_list, count
    if isinstance(value, dict):
        redacted_dict = {}
        count = 0
        for key, item in value.items():
            redacted_item, replacements = redact_public_value(item, personal_values)
            redacted_dict[key] = redacted_item
            count += replacements
        return redacted_dict, count
    return value, 0


def read_supplemental_parts(archive: zipfile.ZipFile) -> list[dict]:
    names = [
        name
        for name in archive.namelist()
        if re.fullmatch(r"word/(?:header|footer)\d+\.xml", name)
    ]
    parts = []
    for name in sorted(names, key=natural_key):
        root = ET.fromstring(archive.read(name))
        paragraphs = []
        for paragraph in root.iter(f"{{{W_NS}}}p"):
            paragraphs.append(
                "".join((node.text or "") for node in paragraph.iter(f"{{{W_NS}}}t"))
            )
        parts.append({"part": name, "text": "\n".join(paragraphs), "paragraphs": paragraphs})
    return parts


def source_paragraphs(archive: zipfile.ZipFile) -> list[str]:
    root = ET.fromstring(archive.read("word/document.xml"))
    body = root.find(f".//{{{W_NS}}}body")
    if body is None:
        return []
    return [
        "".join((node.text or "") for node in paragraph.iter(f"{{{W_NS}}}t"))
        for paragraph in body.iter(f"{{{W_NS}}}p")
    ]


def iter_blocks(parent):
    """Percorre parágrafos e tabelas na ordem real do documento."""
    if isinstance(parent, DocxDocument):
        parent_element = parent.element.body
    elif isinstance(parent, _Cell):
        parent_element = parent._tc
    else:
        raise ValueError("Elemento DOCX inesperado.")
    for child in parent_element.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, parent)
        elif child.tag == qn("w:tbl"):
            yield Table(child, parent)


def heading_level(paragraph: Paragraph):
    style_name = paragraph.style.name or ""
    match = re.match(r"^(?:Heading|Título)\s*(\d+)", style_name, re.IGNORECASE)
    return int(match.group(1)) if match else None


def is_list(paragraph: Paragraph) -> bool:
    properties = paragraph._p.find(qn("w:pPr"))
    return properties is not None and properties.find(qn("w:numPr")) is not None


def paragraph_payload(paragraph: Paragraph, paragraph_id: str) -> dict:
    level = heading_level(paragraph)
    listed = is_list(paragraph)
    payload = {
        "id": paragraph_id,
        "semanticType": "heading" if level else ("list-item" if listed else "paragraph"),
        "text": (paragraph.text or "").strip(),
    }
    if level:
        payload["headingLevel"] = level
    return payload


def parse_heading(text: str) -> tuple[str, str]:
    numbered = re.match(r"^((?:\d+\.)*\d+)\s+(.*)$", text)
    if numbered:
        return numbered.group(1), numbered.group(2).strip()
    annex = re.match(r"^(Anexo\s+[A-Z])\s*[-–]?\s*(.*)$", text, re.IGNORECASE)
    if annex:
        return annex.group(1), annex.group(2).strip()
    annex_child = re.match(r"^([A-Z]\.\d+(?:\.\d+)*)\s+(.*)$", text)
    if annex_child:
        return annex_child.group(1), annex_child.group(2).strip()
    return "", text


def is_navigation_root(number: str, title: str) -> bool:
    return not number and norm(title) in NAVIGATION_ROOTS


def extract_operational_metadata(document: DocxDocument) -> dict:
    text = "\n".join((paragraph.text or "") for paragraph in document.paragraphs)
    version_match = re.search(r"\bVersão\s*:\s*([0-9]+(?:\.[0-9]+)+)", text, re.IGNORECASE)
    date_match = re.search(r"\bData\s*:\s*([^\r\n]+)", text, re.IGNORECASE)
    nature_match = re.search(r"\bNatureza\s*:\s*([^\r\n]+)", text, re.IGNORECASE)
    code_match = re.search(r"\bPOP-[A-Z]+-[A-Z]+-\d+\b", text)
    return {
        "documentCode": code_match.group(0) if code_match else "POP-DLE-HID-001",
        "version": version_match.group(1) if version_match else "",
        "dateLabel": date_match.group(1).strip() if date_match else "",
        "nature": nature_match.group(1).strip() if nature_match else "",
        "jurisdiction": "Estado do Paraná",
        "organization": "Instituto Água e Terra - IAT",
        "validationStatus": "minuta técnica pendente de validação humana e institucional",
    }


def build_id_maps(previous: dict) -> tuple[dict, dict, dict, int]:
    by_number_title: dict[tuple[str, str], str] = {}
    by_title: dict[str, str] = {}
    by_number: dict[str, str] = {}
    maximum = 0
    for section in previous.get("sections", []):
        section_id = section["id"]
        number = section.get("number") or ""
        title = norm(section.get("title") or "")
        by_number_title[(number, title)] = section_id
        by_title.setdefault(title, section_id)
        if number:
            by_number.setdefault(number, section_id)
        match = re.search(r"(\d+)$", section_id)
        if match:
            maximum = max(maximum, int(match.group(1)))
    return by_number_title, by_title, by_number, maximum


def extract_assets(document: DocxDocument) -> tuple[list[dict], dict[str, dict], dict[str, bytes]]:
    image_relationships = [
        (relationship_id, relationship)
        for relationship_id, relationship in document.part.rels.items()
        if "image" in relationship.reltype
    ]
    unique_parts: dict[str, object] = {}
    for _, relationship in image_relationships:
        part_name = str(relationship.target_part.partname).lstrip("/")
        unique_parts.setdefault(part_name, relationship.target_part)

    assets: list[dict] = []
    by_part: dict[str, dict] = {}
    payloads: dict[str, bytes] = {}
    for index, part_name in enumerate(sorted(unique_parts, key=natural_key), 1):
        part = unique_parts[part_name]
        payload = part.blob
        extension = Path(part_name).suffix.lower() or mimetypes.guess_extension(part.content_type) or ".bin"
        file_name = f"pop-image-{index:03d}{extension}"
        image = getattr(part, "image", None)
        asset = {
            "id": f"pop-asset-{index:03d}",
            "documentId": "pop",
            "originalPart": part_name,
            "originalFileName": Path(part_name).name,
            "fileName": file_name,
            "publicPath": f"/source-assets/{file_name}",
            "mimeType": getattr(part, "content_type", None) or mimetypes.guess_type(file_name)[0],
            "bytes": len(payload),
            "widthPx": getattr(image, "px_width", None),
            "heightPx": getattr(image, "px_height", None),
            "sha256": sha256_bytes(payload),
        }
        assets.append(asset)
        by_part[part_name] = asset
        payloads[file_name] = payload

    by_relationship: dict[str, dict] = {}
    for relationship_id, relationship in image_relationships:
        part_name = str(relationship.target_part.partname).lstrip("/")
        by_relationship[relationship_id] = by_part[part_name]
    return assets, by_relationship, payloads


def extract_pop(source: Path, source_bytes: bytes, previous: dict) -> tuple[dict, dict[str, bytes], dict]:
    document = docx.Document(io.BytesIO(source_bytes))
    archive = zipfile.ZipFile(io.BytesIO(source_bytes))
    try:
        source_nodes = source_paragraphs(archive)
        raw_core = read_zip_properties(archive, "docProps/core.xml")
        raw_application = read_zip_properties(archive, "docProps/app.xml")
        redaction_values = personal_metadata_values(raw_core, raw_application)
        core = public_properties(
            raw_core,
            PUBLIC_CORE_PROPERTIES,
        )
        application = public_properties(
            raw_application,
            PUBLIC_APPLICATION_PROPERTIES,
        )
        # Propriedades customizadas podem conter nomes, caminhos ou identificadores
        # arbitrários e não são necessárias para a experiência pública.
        custom = {}
        supplemental = read_supplemental_parts(archive)
    finally:
        archive.close()

    by_number_title, by_title, by_number, maximum_section = build_id_maps(previous)
    used_ids: set[str] = set()

    def next_section_id(number: str, title: str) -> tuple[str, bool]:
        nonlocal maximum_section
        candidates = (
            by_number_title.get((number, norm(title))),
            by_title.get(norm(title)),
            by_number.get(number) if number else None,
        )
        for section_id in candidates:
            if section_id and section_id not in used_ids:
                used_ids.add(section_id)
                return section_id, False
        maximum_section += 1
        section_id = f"pop-section-{maximum_section:03d}"
        used_ids.add(section_id)
        return section_id, True

    assets, assets_by_relationship, asset_payloads = extract_assets(document)
    blocks: list[dict] = []
    sections: list[dict] = []
    tables: list[dict] = []
    figures: list[dict] = []
    extracted_nodes: list[str] = []
    current_section = None
    navigation = False
    navigation_level = 0
    heading_count = 0
    list_item_count = 0
    source_body_paragraph_count = 0
    table_paragraph_count = 0
    new_sections: list[str] = []

    for item in iter_blocks(document):
        if isinstance(item, Paragraph):
            source_body_paragraph_count += 1
            extracted_nodes.append(item.text or "")
            text = (item.text or "").strip()
            level = heading_level(item)
            block_id = f"pop-block-{len(blocks) + 1:04d}"
            if level:
                heading_count += 1
                number, title = parse_heading(text)
                section_id, is_new = next_section_id(number, title)
                if is_new:
                    new_sections.append(f"{number} {title}".strip())
                if is_navigation_root(number, title):
                    navigation, navigation_level = True, level
                elif navigation and level <= navigation_level:
                    navigation = False
                parent_id = None
                for prior in reversed(sections):
                    if prior["level"] < level:
                        parent_id = prior["id"]
                        break
                current_section = {
                    "id": section_id,
                    "level": level,
                    "number": number,
                    "title": title,
                    "fullTitle": text,
                    "headingBlockId": block_id,
                    "parentId": parent_id,
                    "navigationOnly": navigation,
                    "blockIds": [],
                }
                sections.append(current_section)
            drawings = item._p.findall(".//" + qn("w:drawing"))
            if not text and not drawings:
                continue
            block = {
                "id": block_id,
                "sourceIndex": len(blocks) + 1,
                "type": "paragraph",
                "navigationOnly": navigation,
                "sectionId": current_section["id"] if current_section else None,
                "paragraph": paragraph_payload(item, f"pop-paragraph-{len(blocks) + 1:04d}"),
            }
            if block["paragraph"]["semanticType"] == "list-item":
                list_item_count += 1
            blocks.append(block)
            if current_section and block_id != current_section["headingBlockId"]:
                current_section["blockIds"].append(block_id)

            for drawing in item._p.findall(".//" + qn("a:blip")):
                relationship_id = drawing.get(qn("r:embed"))
                asset = assets_by_relationship.get(relationship_id)
                if not asset:
                    continue
                figure_number = len(figures) + 1
                figures.append(
                    {
                        "id": f"pop-figure-{figure_number:03d}",
                        "number": figure_number,
                        "title": text or f"Figura {figure_number}",
                        "caption": text or f"Figura {figure_number}",
                        "blockId": block_id,
                        "assetId": asset["id"],
                        "publicPath": asset["publicPath"],
                        "altText": text or f"Figura {figure_number} do POP",
                    }
                )
            continue

        table_number = len(tables) + 1
        table_id = f"pop-table-{table_number:03d}"
        rows = []
        for row_index, row in enumerate(item.rows, 1):
            cells = []
            for cell_index, cell in enumerate(row.cells, 1):
                paragraphs = []
                for paragraph_index, paragraph in enumerate(cell.paragraphs, 1):
                    table_paragraph_count += 1
                    extracted_nodes.append(paragraph.text or "")
                    paragraphs.append({"index": paragraph_index, "text": paragraph.text or ""})
                cells.append(
                    {
                        "index": cell_index,
                        "text": (cell.text or "").strip(),
                        "paragraphs": paragraphs,
                    }
                )
            rows.append({"index": row_index, "isHeader": row_index == 1, "cells": cells})

        caption = ""
        for prior in reversed(blocks[-4:]):
            prior_text = prior.get("paragraph", {}).get("text", "")
            if re.match(r"^(Quadro|Tabela)\s+\d+", prior_text, re.IGNORECASE):
                caption = prior_text
                break
        caption_match = re.match(
            r"^(Quadro|Tabela)\s+(\d+)\s*[-–]\s*(.*)$", caption, re.IGNORECASE
        )
        label_type = caption_match.group(1).title() if caption_match else "Quadro"
        tables.append(
            {
                "id": table_id,
                "sourceTableIndex": table_number,
                "caption": caption or f"Quadro {table_number}",
                "labelType": label_type,
                "labelNumber": int(caption_match.group(2)) if caption_match else table_number,
                "title": (
                    caption_match.group(3).strip()
                    if caption_match
                    else (caption or f"Quadro {table_number}")
                ),
                "navigationOnly": navigation,
                "rowCount": len(rows),
                "columnCount": len(rows[0]["cells"]) if rows else 0,
                "rows": rows,
            }
        )
        block_id = f"pop-block-{len(blocks) + 1:04d}"
        blocks.append(
            {
                "id": block_id,
                "sourceIndex": len(blocks) + 1,
                "type": "table",
                "navigationOnly": navigation,
                "sectionId": current_section["id"] if current_section else None,
                "tableId": table_id,
                "caption": caption,
            }
        )
        if current_section:
            current_section["blockIds"].append(block_id)

    mismatch_indexes = [
        index
        for index, (source_text, extracted_text) in enumerate(
            zip(source_nodes, extracted_nodes, strict=False)
        )
        if normalized_fidelity_text(source_text) != normalized_fidelity_text(extracted_text)
    ]
    fidelity = {
        "sourceParagraphCount": len(source_nodes),
        "extractedParagraphCount": len(extracted_nodes),
        "normalizedTextMismatchCount": len(mismatch_indexes),
        "firstMismatchIndex": mismatch_indexes[0] if mismatch_indexes else None,
        "pass": len(source_nodes) == len(extracted_nodes) and not mismatch_indexes,
    }

    source_stat = source.stat()
    source_info = {
        "fileName": source.name,
        "bytes": len(source_bytes),
        "lastModifiedUtc": iso_utc(source_stat.st_mtime),
        "sha256": sha256_bytes(source_bytes),
    }
    generated_at = source_info["lastModifiedUtc"]
    operational = extract_operational_metadata(document)
    core_version_match = re.search(
        r"\b(?:V|Versão\s*)([0-9]+(?:\.[0-9]+)+)",
        f"{core.get('title', '')} {core.get('description', '')}",
        re.IGNORECASE,
    )
    core_properties_version = core_version_match.group(1) if core_version_match else None
    learning_sections = [section for section in sections if not section["navigationOnly"]]
    navigation_sections = [section for section in sections if section["navigationOnly"]]
    included_blocks = [block["id"] for block in blocks if not block["navigationOnly"]]
    excluded_blocks = [block["id"] for block in blocks if block["navigationOnly"]]
    stored_body_paragraph_count = sum(block["type"] == "paragraph" for block in blocks)
    searchable_paragraph_nodes = stored_body_paragraph_count + table_paragraph_count
    output = {
        "schemaVersion": "1.0.0",
        "generatedAt": generated_at,
        "id": "pop",
        "kind": "standard-operating-procedure",
        "title": "POP de Licenciamento Ambiental de Empreendimentos Hidrelétricos no IAT",
        "source": source_info,
        "metadata": {
            "core": core,
            "application": application,
            "custom": custom,
            "operational": operational,
            "provenance": {
                "pipeline": "tools/extract_pop.py",
                "pipelineVersion": PIPELINE_VERSION,
                "sourceIdentityStatus": "validada por nome, tamanho e SHA-256",
                "contentStatus": "minuta técnica pendente de validação humana e institucional",
                "coverVersionIsOperationalVersion": True,
                "operationalVersionAuthority": "texto visível da capa do documento",
                "corePropertiesVersion": core_properties_version,
                "corePropertiesStatus": (
                    "metadado interno do Word desatualizado; preservado para auditoria e "
                    "não usado como versão operacional"
                    if core_properties_version and core_properties_version != operational["version"]
                    else "coerente com a versão operacional"
                ),
            },
            "supplementalParts": supplemental,
        },
        "assets": assets,
        "blocks": blocks,
        "sections": sections,
        "tables": tables,
        "figures": figures,
        "flowcharts": previous.get("flowcharts", []),
        "learningContent": {
            "includedBlockIds": included_blocks,
            "excludedNavigationBlockIds": excluded_blocks,
            "substantiveSectionIds": [section["id"] for section in learning_sections],
            "navigationSectionIds": [section["id"] for section in navigation_sections],
        },
        "stats": {
            "bodyBlockCount": len(blocks),
            "bodyParagraphCount": stored_body_paragraph_count,
            "sourceBodyParagraphCount": source_body_paragraph_count,
            "tableParagraphCount": table_paragraph_count,
            # Mantém a métrica histórica usada pela interface: parágrafos
            # efetivamente armazenados/pesquisáveis, sem os 26 vazios do corpo.
            "allDocumentParagraphNodes": searchable_paragraph_nodes,
            "searchableParagraphNodes": searchable_paragraph_nodes,
            "sourceParagraphNodeCount": len(source_nodes),
            "tableCount": len(tables),
            "imageAssetCount": len(assets),
            "figureCount": len(figures),
            "headingCount": heading_count,
            "listItemCount": list_item_count,
            "sectionCount": len(sections),
            "learningSectionCount": len(learning_sections),
            "navigationSectionCount": len(navigation_sections),
            "navigationOnlyBlockCount": len(excluded_blocks),
            "substantiveBlockCount": len(included_blocks),
        },
    }
    output, privacy_redaction_count = redact_public_value(output, redaction_values)
    output["metadata"]["provenance"]["privacyRedactionCount"] = privacy_redaction_count
    output["metadata"]["provenance"]["privacyRedactionPolicy"] = (
        "nomes presentes em metadados de autoria são substituídos na cópia pública"
    )
    diagnostics = {
        "fidelity": fidelity,
        "newSections": new_sections,
        "preservedSectionIds": len(used_ids) - len(new_sections),
        "privacyRedactionCount": privacy_redaction_count,
    }
    return output, asset_payloads, diagnostics


def build_catalog(pop: dict, flows: dict) -> dict:
    return {
        "schemaVersion": "1.1.0",
        "generatedAt": pop["generatedAt"],
        "documents": [
            {
                "id": "pop",
                "title": pop["title"],
                "dataPath": "./pop-content.json",
                "kind": pop["kind"],
                "source": pop["source"],
                "sections": len(pop["sections"]),
                "learningSections": pop["stats"]["learningSectionCount"],
                "navigationSections": pop["stats"]["navigationSectionCount"],
                "tables": len(pop["tables"]),
                "figures": len(pop["figures"]),
            },
            {
                "id": "fluxogramas",
                "title": flows["title"],
                "dataPath": "./flowcharts-content.json",
                "kind": flows["kind"],
                "source": {
                    key: flows.get("source", {}).get(key)
                    for key in ("fileName", "bytes", "lastModifiedUtc", "sha256")
                    if flows.get("source", {}).get(key) is not None
                },
                "flowcharts": len(flows.get("flowcharts", [])),
                "variants": ["original", "simplificado", "completo"],
            },
        ],
        "assetsManifestPath": "/source-assets/asset-manifest.json",
        "suggestedLearningModules": [
            {
                "id": "fundamentos",
                "sectionNumbers": ["1", "2", "3", "4", "5", "6", "7"],
                "title": "Fundamentos, normas e método de análise",
            },
            {
                "id": "enquadramento",
                "sectionNumbers": ["8"],
                "title": "Enquadramento por tipologia, modalidade e estudo",
            },
            {
                "id": "modalidades",
                "sectionNumbers": ["9", "10", "11", "12", "13", "14", "15", "16", "17"],
                "title": "Modalidades, fases e situações especiais",
            },
            {
                "id": "documentos-pacuera",
                "sectionNumbers": ["18"],
                "title": "Documentos técnicos, estudos ambientais e PACUERA",
            },
            {
                "id": "analise-tecnica",
                "sectionNumbers": ["19", "20", "21", "22", "23", "24", "25"],
                "title": "Análise técnica, vistoria, suficiência e condicionantes",
            },
            {
                "id": "qualidade",
                "sectionNumbers": [
                    "26",
                    "27",
                    "Anexo A",
                    "Anexo B",
                    "Anexo C",
                    "Anexo D",
                    "Anexo E",
                    "Anexo F",
                    "Referências",
                ],
                "title": "Produtos técnicos, qualidade, rastreabilidade e anexos",
            },
        ],
    }


def build_asset_manifest(pop: dict, flows: dict) -> dict:
    flow_assets = flows.get("assets", [])
    return {
        "schemaVersion": "1.1.0",
        "generatedAt": pop["generatedAt"],
        "assetCount": len(pop["assets"]) + len(flow_assets),
        "documents": [
            {
                "id": "pop",
                "source": pop["source"],
                "assetCount": len(pop["assets"]),
            },
            {
                "id": "fluxogramas",
                "source": {
                    key: flows.get("source", {}).get(key)
                    for key in ("fileName", "bytes", "lastModifiedUtc", "sha256")
                    if flows.get("source", {}).get(key) is not None
                },
                "assetCount": len(flow_assets),
            },
        ],
        "assets": [*pop["assets"], *flow_assets],
    }


def make_check(check_id: str, expected, actual, details=None) -> dict:
    check = {"id": check_id, "expected": expected, "actual": actual, "pass": expected == actual}
    if details is not None:
        check["details"] = details
    return check


def build_validation(
    pop: dict,
    flows: dict,
    manifest: dict,
    diagnostics: dict,
    pop_asset_payloads: dict[str, bytes],
) -> dict:
    section_263 = next((section for section in pop["sections"] if section.get("number") == "26.3"), None)
    quadros = sum(table.get("labelType") == "Quadro" for table in pop["tables"])
    tabelas = sum(table.get("labelType") == "Tabela" for table in pop["tables"])
    flow_assets = flows.get("assets", [])

    all_flow_assets_valid = True
    for asset in flow_assets:
        path = ASSET_DIR / asset["fileName"]
        if not path.is_file() or sha256_bytes(path.read_bytes()) != asset.get("sha256"):
            all_flow_assets_valid = False
            break
    all_pop_assets_valid = all(
        asset["fileName"] in pop_asset_payloads
        and sha256_bytes(pop_asset_payloads[asset["fileName"]]) == asset["sha256"]
        and len(pop_asset_payloads[asset["fileName"]]) == asset["bytes"]
        for asset in pop["assets"]
    )
    all_dimensions_present = all(
        asset.get("widthPx") and asset.get("heightPx") for asset in manifest["assets"]
    )

    checks = [
        make_check("pop-source-file-name", EXPECTED["fileName"], pop["source"]["fileName"]),
        make_check("pop-source-bytes", EXPECTED["bytes"], pop["source"]["bytes"]),
        make_check("pop-source-sha256", EXPECTED["sha256"], pop["source"]["sha256"]),
        make_check(
            "pop-operational-version",
            EXPECTED["version"],
            pop["metadata"]["operational"]["version"],
        ),
        make_check(
            "pop-operational-version-authority",
            "texto visível da capa do documento",
            pop["metadata"]["provenance"]["operationalVersionAuthority"],
        ),
        make_check(
            "pop-stale-core-version-explicitly-governed",
            True,
            (
                pop["metadata"]["provenance"]["corePropertiesVersion"] == "1.2"
                and "desatualizado" in pop["metadata"]["provenance"]["corePropertiesStatus"]
            ),
        ),
        make_check(
            "pop-public-metadata-sanitized",
            True,
            all(
                key not in pop["metadata"]["core"]
                for key in ("creator", "lastModifiedBy")
            )
            and all(
                key not in pop["metadata"]["application"]
                for key in ("Manager", "Company", "HyperlinkBase")
            )
            and pop["metadata"]["custom"] == {},
        ),
        make_check(
            "pop-personal-author-redactions-applied",
            True,
            diagnostics.get("privacyRedactionCount", 0) > 0,
            {
                "count": diagnostics.get("privacyRedactionCount", 0),
                "replacement": "[nome removido por privacidade]",
            },
        ),
        make_check("pop-section-count", EXPECTED["sections"], len(pop["sections"])),
        make_check(
            "pop-learning-section-count",
            EXPECTED["learningSections"],
            pop["stats"]["learningSectionCount"],
        ),
        make_check(
            "pop-navigation-section-count",
            EXPECTED["navigationSections"],
            pop["stats"]["navigationSectionCount"],
        ),
        make_check("pop-section-26.3-id", "pop-section-102", section_263.get("id") if section_263 else None),
        make_check(
            "pop-section-26.3-substantive",
            True,
            bool(section_263 and not section_263.get("navigationOnly")),
        ),
        make_check(
            "pop-section-26.3-in-learning-content",
            True,
            bool(section_263 and section_263["id"] in pop["learningContent"]["substantiveSectionIds"]),
        ),
        make_check("pop-table-count", EXPECTED["tables"], len(pop["tables"])),
        make_check(
            "pop-table-labels",
            f'{EXPECTED["quadros"]} quadros + {EXPECTED["tabelas"]} tabelas',
            f"{quadros} quadros + {tabelas} tabelas",
        ),
        make_check("pop-image-count", EXPECTED["assets"], len(pop["assets"])),
        make_check("pop-figure-count", EXPECTED["figures"], len(pop["figures"])),
        make_check(
            "pop-paragraph-node-count",
            EXPECTED["paragraphNodes"],
            pop["stats"]["allDocumentParagraphNodes"],
        ),
        make_check("pop-body-block-count", EXPECTED["bodyBlocks"], len(pop["blocks"])),
        make_check(
            "pop-text-fidelity",
            True,
            diagnostics["fidelity"]["pass"],
            diagnostics["fidelity"],
        ),
        make_check("flow-image-count", 21, len(flow_assets)),
        make_check("flowchart-variant-count", 21, len(flows.get("flowcharts", []))),
        make_check("asset-manifest-count", 35, manifest["assetCount"]),
        make_check("all-pop-assets-match-extraction", True, all_pop_assets_valid),
        make_check("all-flow-assets-exist-and-match-manifest", True, all_flow_assets_valid),
        make_check("all-assets-have-dimensions", True, all_dimensions_present),
        make_check(
            "all-tables-have-rows",
            True,
            all(table.get("rowCount", 0) >= 1 for table in pop["tables"]),
        ),
    ]
    return {
        "schemaVersion": "1.1.0",
        "generatedAt": pop["generatedAt"],
        "passed": all(check["pass"] for check in checks),
        "sourceStatus": "minuta técnica pendente de validação humana e institucional",
        "checks": checks,
        "summary": {
            "pop": pop["stats"],
            "fluxogramas": flows.get("stats", {}),
            "totalAssets": manifest["assetCount"],
            "notes": [
                "A identidade da minuta-fonte v1.7 é travada por nome, tamanho e SHA-256.",
                "As seis seções exclusivamente navegacionais são preservadas, mas não viram aulas.",
                "A seção 26.3 é conteúdo substantivo e integra a trilha de aprendizagem.",
                "Os artefatos são propostas de treinamento pendentes de validação humana e institucional.",
            ],
        },
    }


def json_bytes(value: object) -> bytes:
    return (json.dumps(value, ensure_ascii=False, separators=(",", ":")) + "\n").encode("utf-8")


def write_if_changed(path: Path, payload: bytes) -> bool:
    if path.is_file() and path.read_bytes() == payload:
        return False
    path.parent.mkdir(parents=True, exist_ok=True)
    temporary = path.with_suffix(path.suffix + ".tmp")
    temporary.write_bytes(payload)
    temporary.replace(path)
    return True


def publish(
    pop: dict,
    catalog: dict,
    manifest: dict,
    validation: dict,
    asset_payloads: dict[str, bytes],
) -> list[str]:
    changed = []
    for file_name, payload in asset_payloads.items():
        if write_if_changed(ASSET_DIR / file_name, payload):
            changed.append(f"public/source-assets/{file_name}")

    expected_pop_assets = set(asset_payloads)
    for stale_path in ASSET_DIR.glob("pop-image-*"):
        if stale_path.name not in expected_pop_assets:
            stale_path.unlink()
            changed.append(f"removido: public/source-assets/{stale_path.name}")

    artifacts = {
        OUT_JSON: json_bytes(pop),
        CATALOG_JSON: json_bytes(catalog),
        ASSET_MANIFEST: json_bytes(manifest),
        VALIDATION_JSON: json_bytes(validation),
    }
    for path, payload in artifacts.items():
        if write_if_changed(path, payload):
            changed.append(path.relative_to(ROOT).as_posix())
    return changed


def parse_args(argv: list[str]) -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument(
        "source",
        nargs="?",
        type=Path,
        default=DEFAULT_POP_SOURCE,
        help="Caminho da minuta DOCX v1.7.",
    )
    parser.add_argument(
        "--check-only",
        action="store_true",
        help="Executa extração e gates em memória, sem alterar os artefatos.",
    )
    return parser.parse_args(argv)


def main(argv: list[str] | None = None) -> int:
    args = parse_args(argv or sys.argv[1:])
    if args.source is None:
        print(
            "FALHA: informe o DOCX por argumento ou pela variável IAT_POP_SOURCE.",
            file=sys.stderr,
        )
        return 2
    source = args.source.expanduser().resolve()
    if not source.is_file():
        print(f"FALHA: fonte não localizada: {source}", file=sys.stderr)
        return 2
    if not FLOW_JSON.is_file():
        print(f"FALHA: fluxogramas não localizados: {FLOW_JSON}", file=sys.stderr)
        return 2

    source_bytes = source.read_bytes()
    identity_errors = []
    if source.name != EXPECTED["fileName"]:
        identity_errors.append(f"nome: esperado {EXPECTED['fileName']!r}, recebido {source.name!r}")
    if len(source_bytes) != EXPECTED["bytes"]:
        identity_errors.append(f"tamanho: esperado {EXPECTED['bytes']}, recebido {len(source_bytes)}")
    source_sha = sha256_bytes(source_bytes)
    if source_sha != EXPECTED["sha256"]:
        identity_errors.append(f"SHA-256: esperado {EXPECTED['sha256']}, recebido {source_sha}")
    if identity_errors:
        print("FALHA: a fonte não corresponde à minuta v1.7 aprovada para extração.", file=sys.stderr)
        for error in identity_errors:
            print(f"  - {error}", file=sys.stderr)
        return 1

    previous = json.loads(OUT_JSON.read_text(encoding="utf-8")) if OUT_JSON.is_file() else {}
    flows = json.loads(FLOW_JSON.read_text(encoding="utf-8"))
    pop, asset_payloads, diagnostics = extract_pop(source, source_bytes, previous)
    catalog = build_catalog(pop, flows)
    manifest = build_asset_manifest(pop, flows)
    validation = build_validation(pop, flows, manifest, diagnostics, asset_payloads)

    print(
        "POP v{version}: {sections} seções ({learning} substantivas), "
        "{tables} tabelas, {figures} figuras, {nodes} parágrafos.".format(
            version=pop["metadata"]["operational"]["version"],
            sections=len(pop["sections"]),
            learning=pop["stats"]["learningSectionCount"],
            tables=len(pop["tables"]),
            figures=len(pop["figures"]),
            nodes=pop["stats"]["allDocumentParagraphNodes"],
        )
    )
    print(
        f"SHA-256: {pop['source']['sha256']} | IDs preservados: "
        f"{diagnostics['preservedSectionIds']} | IDs novos: {len(diagnostics['newSections'])}"
    )
    failed = [check for check in validation["checks"] if not check["pass"]]
    if failed:
        print(f"FALHA: {len(failed)} gate(s) de extração não passaram.", file=sys.stderr)
        for check in failed:
            print(
                f"  - {check['id']}: esperado {check['expected']!r}, recebido {check['actual']!r}",
                file=sys.stderr,
            )
        print("Nenhum artefato foi publicado.", file=sys.stderr)
        return 1

    if args.check_only:
        print("OK: extração validada em memória; nenhum arquivo foi alterado.")
        return 0

    changed = publish(pop, catalog, manifest, validation, asset_payloads)
    if changed:
        print("Artefatos atualizados:")
        for item in changed:
            print(f"  - {item}")
    else:
        print("OK: cadeia idempotente; os artefatos já estavam atualizados.")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
